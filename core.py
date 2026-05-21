"""
Core scraping, filtering, ranking, and rendering logic.
Shared by the Streamlit app (app.py) and the email cron worker
(.github/workflows/send_emails.yml -> scripts/send_emails.py).
"""

import hashlib
import json
import re
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------- Banneker palette ----------

NAVY = "#181F64"
DEEP_NAVY = "#0D0537"
DARK_GREY = "#59595B"
PERIWINKLE = "#76A3E3"
ICE_BLUE = "#E4EDF9"
MID_GREY = "#A6A6A6"
BODY_GREY = "#3A3A3C"

# ---------- News filter ----------

NON_NEWS_PATTERNS = [
    r"\bwebinar\b",
    r"\bwhitepaper\b",
    r"\bwhite\s+paper\b",
    r"\b\d+\s+(things|ways|tips|steps|reasons|key|top|best)\b",
    r"\bmust\s+know\b",
    r"\bguide\s+to\b",
    r"\bhow\s+to\b",
    r"\b(beginner'?s|complete|ultimate|essential|definitive|practical|comprehensive)\s+guide\b",
    r"\bbest\s+practices\b",
    r"\bchecklist\b",
    r"\bpodcast\b",
    r"\bep(isode)?\.?\s*\d+\b",
    r"\bspotlight(?:s|ing)?\b",
    r"\binterview\b",
    r"\bq\s*&\s*a\b",
    r"\bama\b",
    r"\bsponsored\b",
    r"\bexplained\b",
    r"\beverything\s+you\s+(need\s+to\s+)?know\b",
    r"\bwhat\s+is\b",
    r"\bcase\s+study\b",
    r"\bcareer\s+spotlight\b",
    r"\bopinion:\s*",
    r"\bcommentary:\s*",
    r"\bcolumn:\s*",
]

NEWS_SIGNAL_WORDS = [
    "breach", "breached", "attack", "attacks", "attacked", "hacked", "hack",
    "ransomware", "malware", "exploit", "exploited", "zero-day", "zero day",
    "fined", "fine", "penalty", "settles", "settlement", "indicted", "arrested",
    "acquires", "acquired", "acquisition", "buys", "merger", "merges",
    "raises", "raised", "funding", "ipo", "ipos",
    "launches", "unveils", "announces", "appoints", "names", "hires",
    "shuts down", "shut down", "halts", "halted", "outage", "disrupts",
    "warns", "warning", "alert", "advisory", "vulnerability", "cve",
    "passes", "passed", "enacts", "enacted", "approves", "approved", "directive",
    "investigation", "probe", "indictment", "lawsuit", "sued",
    "leaked", "leak", "exposed", "stolen", "theft",
    "approval", "cleared", "phase 3", "phase 2", "primary endpoint",
]

LOW_QUALITY_SOURCES = [
    "openpr", "pr newswire", "businesswire", "globe newswire",
    "globenewswire", "press release", "marketwatch sponsored",
    "imdb", "yahoo finance", "the joplin globe",
]


def is_news_story(title: str) -> bool:
    low = title.lower()
    for pat in NON_NEWS_PATTERNS:
        if re.search(pat, low):
            return False
    return True


def news_score(item: dict, theme_idx: int, total_themes: int, lookback_days: int) -> float:
    # First theme is highest weight (typically Breaches), descending
    theme_weight = 1.0 - 0.08 * theme_idx
    theme_weight = max(0.4, theme_weight)
    score = theme_weight
    days_old = (datetime.now(timezone.utc) - item["published"]).days
    recency = max(0.0, 1.0 - days_old / max(1, lookback_days))
    score *= (0.55 + 0.45 * recency)
    title_low = item["title"].lower()
    signal_hits = sum(1 for w in NEWS_SIGNAL_WORDS if w in title_low)
    score *= (1.0 + 0.15 * min(signal_hits, 4))
    src = (item.get("source") or "").lower()
    if any(s in src for s in LOW_QUALITY_SOURCES):
        score *= 0.4
    return score


def pick_top_stories(grouped: dict, theme_order: list[str], n: int, lookback_days: int) -> list[tuple[dict, str]]:
    scored: list[tuple[dict, str, float]] = []
    for theme_idx, theme in enumerate(theme_order):
        for it in grouped.get(theme, []):
            if not is_news_story(it["title"]):
                continue
            scored.append((it, theme, news_score(it, theme_idx, len(theme_order), lookback_days)))
    scored.sort(key=lambda x: x[2], reverse=True)
    return [(s[0], s[1]) for s in scored[:n]]


# ---------- Scraping ----------

def fetch_rss(query: str) -> list[dict]:
    url = (
        "https://news.google.com/rss/search?"
        f"q={urllib.parse.quote(query)}&hl=en-US&gl=US&ceid=US:en"
    )
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urllib.request.urlopen(req, timeout=20) as r:
            data = r.read()
    except Exception:
        return []

    items = []
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return []

    for item in root.iter("item"):
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        pub_raw = (item.findtext("pubDate") or "").strip()
        source_el = item.find("source")
        source = source_el.text if source_el is not None and source_el.text else ""
        try:
            pub = parsedate_to_datetime(pub_raw)
            if pub.tzinfo is None:
                pub = pub.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        items.append({
            "title": title,
            "link": link,
            "published": pub,
            "source": source,
        })
    return items


def scrape_for_config(config: dict) -> tuple[dict, int]:
    """Run scraping for a single tracker config. Returns (grouped, total)
    where grouped is {theme_name: [items]} in theme order."""
    lookback_days = int(config.get("lookback_days", 7))
    cutoff = datetime.now(timezone.utc) - timedelta(days=lookback_days)
    themes = config.get("themes", [])
    grouped: dict[str, list[dict]] = {}
    seen_links: set[str] = set()
    seen_titles: set[str] = set()

    for theme in themes:
        name = theme.get("name", "").strip()
        if not name:
            continue
        queries = [q.strip() for q in theme.get("queries", []) if q.strip()]
        items = []
        for q in queries:
            for it in fetch_rss(q):
                if it["published"] < cutoff:
                    continue
                if it["link"] in seen_links:
                    continue
                title_key = it["title"].lower()[:120]
                if title_key in seen_titles:
                    continue
                if not is_news_story(it["title"]):
                    continue
                seen_links.add(it["link"])
                seen_titles.add(title_key)
                items.append(it)
        items.sort(key=lambda x: x["published"], reverse=True)
        grouped[name] = items

    total = sum(len(v) for v in grouped.values())
    return grouped, total


def config_hash(config: dict) -> str:
    """Stable hash of the parts of a config that affect scraping. For cache keys."""
    relevant = {
        "themes": config.get("themes", []),
        "lookback_days": config.get("lookback_days", 7),
    }
    return hashlib.md5(json.dumps(relevant, sort_keys=True).encode("utf-8")).hexdigest()[:10]


# ---------- DOCX export ----------

def _set_font(rPr, font_name="Inter", size_pt=11, color=None, bold=False, italic=False, underline=False):
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)


def _style_run(run, size_pt=11, color_hex="59595B", bold=False, italic=False):
    run.font.name = "Inter"
    run.font.size = Pt(size_pt)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Inter")
    rFonts.set(qn("w:hAnsi"), "Inter")


def _add_hyperlink(paragraph, url, text, size_pt=11):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    _set_font(rPr, size_pt=size_pt, color="181F64", underline=True)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx_bytes(grouped: dict, theme_order: list[str], total: int,
                     config: dict) -> bytes:
    lookback_days = int(config.get("lookback_days", 7))
    title = config.get("title") or f"{config.get('name', 'News')} Brief"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.footer_distance = Inches(0.4)

    today_pretty = datetime.now().strftime("%B %d, %Y")

    p = doc.add_paragraph()
    _style_run(p.add_run(title), size_pt=20, color_hex="181F64", bold=True)
    p = doc.add_paragraph()
    _style_run(p.add_run(today_pretty), size_pt=12, color_hex="76A3E3", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _style_run(
        p.add_run(
            f"Past {lookback_days} days  |  {total} stories  |  "
            f"{len(theme_order)} themes  |  Source: Google News RSS"
        ),
        size_pt=9, color_hex="59595B", italic=True,
    )

    for theme in theme_order:
        items = grouped.get(theme, [])
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        _style_run(h.add_run(theme), size_pt=13, color_hex="181F64", bold=True)

        if not items:
            empty = doc.add_paragraph()
            _style_run(empty.add_run("No news in the lookback window"),
                       size_pt=10, color_hex="A6A6A6", italic=True)
            continue

        for it in items:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.18)
            bp.paragraph_format.space_after = Pt(2)
            _style_run(bp.add_run("•  "), size_pt=11, color_hex="181F64", bold=True)
            _style_run(bp.add_run(it["published"].strftime("%b %d") + "  "),
                       size_pt=11, color_hex="181F64", bold=True)
            _add_hyperlink(bp, it["link"], it["title"])
            if it["source"]:
                _style_run(bp.add_run(f"  — {it['source']}"),
                           size_pt=10, color_hex="A6A6A6", italic=True)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = ""
    year = datetime.now().year
    _style_run(
        fp.add_run(f"(C) {year} Banneker Partners, LLC. All Rights Reserved. Confidential."),
        size_pt=8, color_hex="59595B",
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------- Email HTML ----------

def build_email_html(grouped: dict, theme_order: list[str], total: int,
                     config: dict, unsubscribe_url: str | None = None) -> str:
    import html as _html

    title = config.get("title") or f"{config.get('name', 'News')} Brief"
    subtitle = config.get("subtitle", "")
    lookback_days = int(config.get("lookback_days", 7))
    show_top_news = config.get("show_top_news", True)
    today_pretty = datetime.now().strftime("%B %d, %Y")

    top_html = ""
    if show_top_news:
        top_n = int(config.get("top_news_count", 4))
        top_picks = pick_top_stories(grouped, theme_order, n=top_n, lookback_days=lookback_days)
        if top_picks:
            rows = []
            for i, (it, theme) in enumerate(top_picks, start=1):
                date = it["published"].strftime("%b %d")
                src = (
                    f"<span style='color:rgba(255,255,255,0.6); font-style:italic; font-size:12px;'> &mdash; {_html.escape(it['source'])}</span>"
                    if it["source"] else ""
                )
                theme_short = theme.split(" (")[0]
                rows.append(
                    f"<div style='margin-bottom:14px; line-height:1.5;'>"
                    f"<span style='color:{PERIWINKLE}; font-weight:800; font-size:13px; margin-right:8px;'>{i:02d}</span>"
                    f"<span style='color:rgba(255,255,255,0.6); font-size:11px; letter-spacing:0.8px; text-transform:uppercase; font-weight:600;'>{_html.escape(theme_short)} &middot; {date}</span><br>"
                    f"<a href='{_html.escape(it['link'])}' style='color:#FFFFFF; font-size:15.5px; font-weight:600; text-decoration:none; border-bottom:1px dotted rgba(255,255,255,0.45);'>{_html.escape(it['title'])}</a>"
                    f"{src}"
                    f"</div>"
                )
            top_html = f"""
            <div style='margin:22px 0; padding:22px 26px; background:{NAVY}; border-radius:4px;'>
              <div style='color:{PERIWINKLE}; font-size:11px; letter-spacing:1.8px; text-transform:uppercase; font-weight:700; margin-bottom:14px;'>
                This week's top news
              </div>
              {''.join(rows)}
            </div>
            """

    theme_html_parts = []
    for i, theme in enumerate(theme_order):
        items = grouped.get(theme, [])
        if not items:
            body = f"<div style='color:{MID_GREY}; font-style:italic; font-size:13px; margin:8px 0;'>No news in the lookback window.</div>"
        else:
            rows = []
            for it in items:
                date = it["published"].strftime("%b %d")
                src = (
                    f" <span style='color:{MID_GREY}; font-style:italic; font-size:12px;'>&mdash; {_html.escape(it['source'])}</span>"
                    if it["source"] else ""
                )
                rows.append(
                    f"<div style='margin-bottom:10px; line-height:1.45; color:{BODY_GREY}; font-size:14.5px;'>"
                    f"<span style='display:inline-block; min-width:56px; color:{NAVY}; font-weight:700; font-size:13px;'>{date}</span>"
                    f"<a href='{_html.escape(it['link'])}' style='color:{NAVY}; text-decoration:underline;'>{_html.escape(it['title'])}</a>"
                    f"{src}"
                    f"</div>"
                )
            body = "".join(rows)
        theme_html_parts.append(f"""
        <div style='margin:36px 0 14px 0;'>
          <div style='color:{PERIWINKLE}; font-size:11px; font-weight:700; letter-spacing:1.8px; text-transform:uppercase; margin-bottom:2px;'>Section {i+1:02d}</div>
          <h2 style='color:{NAVY}; font-weight:800; font-size:22px; margin:0 0 4px 0; line-height:1.2;'>{_html.escape(theme)}</h2>
          <div style='height:2px; background:{NAVY}; width:48px; margin-top:8px;'></div>
        </div>
        {body}
        """)

    unsubscribe_html = ""
    if unsubscribe_url:
        unsubscribe_html = (
            f"<div style='margin-top:24px; font-size:11px; color:{MID_GREY};'>"
            f"<a href='{_html.escape(unsubscribe_url)}' style='color:{MID_GREY};'>Unsubscribe</a> &middot; "
            f"<a href='{_html.escape(unsubscribe_url.replace('unsubscribe', 'view'))}' style='color:{MID_GREY};'>View in browser</a>"
            f"</div>"
        )

    year = datetime.now().year
    subtitle_html = (
        f"<div style='color:{PERIWINKLE}; font-weight:700; font-size:14px; letter-spacing:1.5px; text-transform:uppercase; margin-bottom:14px;'>{_html.escape(subtitle)}</div>"
        if subtitle else ""
    )

    return f"""<!doctype html>
<html><head><meta charset="utf-8"></head>
<body style="font-family: Inter, Calibri, 'Segoe UI', Arial, sans-serif; color:{BODY_GREY}; background:#FFFFFF; max-width:720px; margin:0 auto; padding:32px 28px 48px; font-size:15px; line-height:1.5;">
  <div style='border-top:4px solid {NAVY}; padding-top:18px; margin-bottom:8px;'></div>
  <h1 style='color:{NAVY}; font-weight:800; font-size:34px; line-height:1.1; margin:0 0 6px 0; letter-spacing:-0.5px;'>{_html.escape(title)}</h1>
  {subtitle_html}
  <div style='margin-top:14px; padding:14px 18px; background:{ICE_BLUE}; border-left:5px solid {NAVY}; border-radius:2px;'>
    <div style='color:{NAVY}; font-weight:800; font-size:20px; line-height:1.1; margin-bottom:4px;'>{today_pretty}</div>
    <div style='color:{DARK_GREY}; font-size:12px; letter-spacing:0.5px; text-transform:uppercase; font-weight:600;'>
      Past {lookback_days} days &middot; {total} stories &middot; {len(theme_order)} themes &middot; Source: Google News RSS
    </div>
  </div>
  {top_html}
  {''.join(theme_html_parts)}
  <div style='margin-top:56px; padding-top:16px; border-top:1px solid {ICE_BLUE}; color:{MID_GREY}; font-size:11px;'>
    (C) {year} Banneker Partners, LLC. All Rights Reserved. Confidential.
  </div>
  {unsubscribe_html}
</body></html>
"""
