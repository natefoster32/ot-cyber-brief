"""
OT Cybersecurity Weekly Brief — Streamlit web app
Deployed at share.streamlit.io. Bryan clicks button → scraper runs → branded
brief renders in-page. Powered by Banneker Partners.
"""

import re
import urllib.request
import urllib.parse
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
from io import BytesIO

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Config ----------

LOOKBACK_DAYS = 7

QUERIES = {
    "Breaches & incidents": [
        "OT cyberattack",
        "ICS ransomware",
        "critical infrastructure breach",
        "utility cyberattack",
        "water utility cyberattack",
        "power grid cyberattack",
        "oil gas pipeline cyberattack",
        "manufacturing ransomware shutdown",
        "factory ransomware",
        "energy sector breach",
        "SCADA attack",
        "industrial ransomware",
    ],
    "Competitive / market": [
        "Industrial Defender OT",
        "Claroty OT security",
        "Dragos OT cybersecurity",
        "Nozomi Networks",
        "Armis OT",
        "TXOne Networks",
    ],
    "EU regulation (NIS2 / CRA / DORA)": [
        "NIS2 directive",
        "Cyber Resilience Act",
        "DORA cybersecurity",
        "ENISA OT",
        "NIS2 enforcement fine",
    ],
    "North American regulation (NERC-CIP / TSA / CISA)": [
        "NERC CIP",
        "NERC CIP compliance",
        "NERC CIP audit penalty",
        "FERC cybersecurity",
        "TSA pipeline cybersecurity directive",
        "CISA OT advisory",
        "CISA critical infrastructure",
        "EPA water cybersecurity",
    ],
    "OT / ICS cybersecurity": [
        "OT cybersecurity",
        "ICS cybersecurity",
        "operational technology security",
        "SCADA security",
        "industrial control system security",
    ],
    "Country-specific (Europe)": [
        "Germany cybersecurity regulation OT",
        "Poland critical infrastructure cybersecurity",
        "UK NIS regulations",
        "Italy NIS2",
        "Austria Cybersicherheit OT",
        "Switzerland critical infrastructure cybersecurity",
        "Turkey energy cybersecurity",
        "Ireland NIS2",
    ],
}

# ---------- News filter ----------
# Drop marketing/educational content. Keep actual news.
NON_NEWS_PATTERNS = [
    r"\bwebinar\b",
    r"\bwhitepaper\b",
    r"\bwhite\s+paper\b",
    r"\b\d+\s+(things|ways|tips|steps|reasons|key|top|best)\b",
    r"\bmust\s+know\b",
    r"\bguide\s+to\b",
    r"\bhow\s+to\b",
    r"\b(beginner'?s|complete|ultimate|essential|definitive|practical|comprehensive)\s+guide\b",
    r"\bbest\s+practices\b",
    r"\bchecklist\b",
    r"\bpodcast\b",
    r"\bep(isode)?\.?\s*\d+\b",
    r"\bspotlight(?:s|ing)?\b",
    r"\binterview\b",
    r"\bq\s*&\s*a\b",
    r"\bama\b",
    r"\bsponsored\b",
    r"\bexplained\b",
    r"\beverything\s+you\s+(need\s+to\s+)?know\b",
    r"\bwhat\s+is\b",
    r"\bcase\s+study\b",
    r"\bcareer\s+spotlight\b",
    r"\bopinion:\s*",
    r"\bcommentary:\s*",
    r"\bcolumn:\s*",
]

NEWS_SIGNAL_WORDS = [
    "breach", "breached", "attack", "attacks", "attacked", "hacked", "hack",
    "ransomware", "malware", "exploit", "exploited", "zero-day", "zero day",
    "fined", "fine", "penalty", "settles", "settlement", "indicted", "arrested",
    "acquires", "acquired", "acquisition", "buys", "merger", "merges",
    "raises", "raised", "funding", "ipo", "ipos",
    "launches", "unveils", "announces", "appoints", "names", "hires",
    "shuts down", "shut down", "halts", "halted", "outage", "disrupts",
    "warns", "warning", "alert", "advisory", "vulnerability", "cve",
    "passes", "passed", "enacts", "enacted", "approves", "approved", "directive",
    "investigation", "probe", "indictment", "lawsuit", "sued",
    "leaked", "leak", "exposed", "stolen", "theft",
]

def is_news_story(title: str) -> bool:
    low = title.lower()
    for pat in NON_NEWS_PATTERNS:
        if re.search(pat, low):
            return False
    return True

# Theme weights for ranking — breaches/competitive are highest signal for Bryan
THEME_WEIGHTS = {
    "Breaches & incidents": 1.00,
    "Competitive / market": 0.92,
    "North American regulation (NERC-CIP / TSA / CISA)": 0.72,
    "EU regulation (NIS2 / CRA / DORA)": 0.68,
    "OT / ICS cybersecurity": 0.55,
    "Country-specific (Europe)": 0.50,
}

LOW_QUALITY_SOURCES = [
    "openpr", "pr newswire", "businesswire", "globe newswire",
    "globenewswire", "press release", "marketwatch sponsored",
    "imdb", "yahoo finance", "the joplin globe",
]

def news_score(item: dict, theme: str) -> float:
    score = THEME_WEIGHTS.get(theme, 0.5)
    days_old = (datetime.now(timezone.utc) - item["published"]).days
    recency = max(0.0, 1.0 - days_old / 7.0)
    score *= (0.55 + 0.45 * recency)
    title_low = item["title"].lower()
    signal_hits = sum(1 for w in NEWS_SIGNAL_WORDS if w in title_low)
    score *= (1.0 + 0.15 * min(signal_hits, 4))
    src = (item.get("source") or "").lower()
    if any(s in src for s in LOW_QUALITY_SOURCES):
        score *= 0.4
    return score

def pick_top_stories(grouped: dict, n: int = 4) -> list[tuple[dict, str]]:
    scored: list[tuple[dict, str, float]] = []
    for theme, items in grouped.items():
        for it in items:
            if not is_news_story(it["title"]):
                continue
            scored.append((it, theme, news_score(it, theme)))
    scored.sort(key=lambda x: x[2], reverse=True)
    return [(s[0], s[1]) for s in scored[:n]]


# Banneker palette
NAVY = "#181F64"
DEEP_NAVY = "#0D0537"
DARK_GREY = "#59595B"
PERIWINKLE = "#76A3E3"
ICE_BLUE = "#E4EDF9"
MID_GREY = "#A6A6A6"

# ---------- Scrape ----------

def fetch_rss(query: str) -> list[dict]:
    url = (
        "https://news.google.com/rss/search?"
        f"q={urllib.parse.quote(query)}&hl=en-US&gl=US&ceid=US:en"
    )
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urllib.request.urlopen(req, timeout=20) as r:
            data = r.read()
    except Exception:
        return []

    items = []
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return []

    for item in root.iter("item"):
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        pub_raw = (item.findtext("pubDate") or "").strip()
        source_el = item.find("source")
        source = source_el.text if source_el is not None and source_el.text else ""
        try:
            pub = parsedate_to_datetime(pub_raw)
            if pub.tzinfo is None:
                pub = pub.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        items.append({
            "title": title,
            "link": link,
            "published": pub,
            "source": source,
        })
    return items


@st.cache_data(ttl=1800, show_spinner=False)
def scrape_all() -> tuple[dict, int]:
    """Run all queries, dedupe, return (grouped, total). Cached 30 min."""
    cutoff = datetime.now(timezone.utc) - timedelta(days=LOOKBACK_DAYS)
    grouped: dict[str, list[dict]] = {}
    seen_links: set[str] = set()
    seen_titles: set[str] = set()

    for theme, queries in QUERIES.items():
        items = []
        for q in queries:
            for it in fetch_rss(q):
                if it["published"] < cutoff:
                    continue
                if it["link"] in seen_links:
                    continue
                title_key = it["title"].lower()[:120]
                if title_key in seen_titles:
                    continue
                seen_links.add(it["link"])
                seen_titles.add(title_key)
                items.append(it)
        items.sort(key=lambda x: x["published"], reverse=True)
        grouped[theme] = items

    total = sum(len(v) for v in grouped.values())
    return grouped, total


# ---------- DOCX builder (for download) ----------

def _set_font(rPr, font_name="Inter", size_pt=11, color=None, bold=False, italic=False, underline=False):
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)


def _style_run(run, size_pt=11, color_hex="59595B", bold=False, italic=False):
    run.font.name = "Inter"
    run.font.size = Pt(size_pt)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Inter")
    rFonts.set(qn("w:hAnsi"), "Inter")


def _add_hyperlink(paragraph, url, text, size_pt=11):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    _set_font(rPr, size_pt=size_pt, color="181F64", underline=True)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx_bytes(grouped: dict, total: int) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.footer_distance = Inches(0.4)

    today_pretty = datetime.now().strftime("%B %d, %Y")

    p = doc.add_paragraph()
    _style_run(p.add_run("OT Cybersecurity Weekly Brief"), size_pt=20, color_hex="181F64", bold=True)
    p = doc.add_paragraph()
    _style_run(p.add_run(today_pretty), size_pt=12, color_hex="76A3E3", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _style_run(
        p.add_run(f"Past {LOOKBACK_DAYS} days  |  {total} stories  |  {len(grouped)} themes  |  Source: Google News RSS"),
        size_pt=9, color_hex="59595B", italic=True,
    )

    for theme, items in grouped.items():
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        _style_run(h.add_run(theme), size_pt=13, color_hex="181F64", bold=True)

        if not items:
            empty = doc.add_paragraph()
            _style_run(empty.add_run("No coverage in the lookback window"),
                       size_pt=10, color_hex="A6A6A6", italic=True)
            continue

        for it in items:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.18)
            bp.paragraph_format.space_after = Pt(2)
            _style_run(bp.add_run("•  "), size_pt=11, color_hex="181F64", bold=True)
            _style_run(bp.add_run(it["published"].strftime("%b %d") + "  "),
                       size_pt=11, color_hex="181F64", bold=True)
            _add_hyperlink(bp, it["link"], it["title"])
            if it["source"]:
                _style_run(bp.add_run(f"  — {it['source']}"),
                           size_pt=10, color_hex="A6A6A6", italic=True)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = ""
    year = datetime.now().year
    _style_run(fp.add_run(f"(C) {year} Banneker Partners, LLC. All Rights Reserved. Confidential."),
               size_pt=8, color_hex="59595B")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------- UI ----------

st.set_page_config(
    page_title="OT Cyber Weekly Brief — Banneker Partners",
    page_icon="🛡️",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown(f"""
<style>
  .stApp {{ background: #FFFFFF; }}
  html, body, [class*="css"], [class*="st-"] {{
    font-family: Inter, Calibri, "Segoe UI", Arial, sans-serif !important;
  }}
  /* Force navy on ALL heading levels and any nested spans Streamlit injects */
  h1, h1 *, h2, h2 *, h3, h3 *,
  [data-testid="stMarkdownContainer"] h1,
  [data-testid="stMarkdownContainer"] h1 *,
  [data-testid="stMarkdownContainer"] h2,
  [data-testid="stMarkdownContainer"] h2 *,
  [data-testid="stMarkdownContainer"] h3,
  [data-testid="stMarkdownContainer"] h3 * {{
    color: {NAVY} !important;
    font-weight: 700 !important;
  }}
  /* Body text — darker for stronger contrast */
  p, li, div, span, label {{ color: #3A3A3C; }}
  .block-container {{ padding-top: 2.5rem; padding-bottom: 4rem; max-width: 800px; }}
  /* Primary CTA */
  .stButton > button {{
    background-color: {NAVY} !important;
    color: #FFFFFF !important;
    font-weight: 700 !important;
    border: none !important;
    padding: 16px 36px !important;
    border-radius: 4px !important;
    font-size: 16px !important;
    width: 100%;
    letter-spacing: 0.3px;
    box-shadow: 0 2px 6px rgba(24, 31, 100, 0.18);
    transition: all 0.15s ease;
  }}
  .stButton > button:hover {{
    background-color: {DEEP_NAVY} !important;
    color: #FFFFFF !important;
    box-shadow: 0 4px 10px rgba(24, 31, 100, 0.25);
    transform: translateY(-1px);
  }}
  .stButton > button p {{ color: #FFFFFF !important; font-weight: 700 !important; }}
  /* Download button */
  .stDownloadButton > button {{
    background-color: #FFFFFF !important;
    color: {NAVY} !important;
    border: 1.5px solid {NAVY} !important;
    font-weight: 700 !important;
    padding: 10px 22px !important;
    border-radius: 4px !important;
  }}
  .stDownloadButton > button:hover {{
    background-color: {ICE_BLUE} !important;
  }}
  .stDownloadButton > button p {{ color: {NAVY} !important; font-weight: 700 !important; }}
  /* Links */
  a {{ color: {NAVY} !important; text-decoration: underline; text-decoration-color: rgba(118, 163, 227, 0.5); }}
  a:hover {{ color: {PERIWINKLE} !important; text-decoration-color: {PERIWINKLE}; }}
  /* Hide Streamlit chrome */
  #MainMenu, footer, header[data-testid="stHeader"] {{ visibility: hidden; }}
</style>
""", unsafe_allow_html=True)

# Header (raw HTML to bypass Streamlit's markdown wrappers)
st.markdown(
    f"""
    <div style='border-top: 4px solid {NAVY}; padding-top: 18px; margin-bottom: 8px;'></div>
    <h1 style='color: {NAVY}; font-weight: 800; font-size: 38px; line-height: 1.1; margin: 0 0 6px 0; letter-spacing: -0.5px;'>
        OT Cybersecurity Weekly Brief
    </h1>
    <div style='color: {PERIWINKLE}; font-weight: 700; font-size: 14px; letter-spacing: 1.5px; text-transform: uppercase; margin-bottom: 14px;'>
        Banneker Partners &middot; Industrial Defender Market Intel
    </div>
    <div style='color: #3A3A3C; font-size: 15px; line-height: 1.5; margin-bottom: 24px; max-width: 640px;'>
        Click the button to pull the last 7 days of OT cyber news across breaches, the
        competitive landscape, EU and NA regulation, and ICS security. Fresh every click.
    </div>
    """,
    unsafe_allow_html=True,
)

# Button
go = st.button("Generate this week's brief", type="primary")

if go or st.session_state.get("has_results"):
    if go or "results" not in st.session_state:
        with st.spinner("Pulling stories from Google News..."):
            grouped, total = scrape_all()
        st.session_state["results"] = (grouped, total)
        st.session_state["has_results"] = True
        st.session_state["generated_at"] = datetime.now()
    else:
        grouped, total = st.session_state["results"]

    generated = st.session_state["generated_at"]

    # Apply news filter and re-build ordered, filtered groups using QUERIES order
    filtered_grouped: dict[str, list[dict]] = {}
    for theme in QUERIES.keys():
        items = grouped.get(theme, [])
        filtered_grouped[theme] = [it for it in items if is_news_story(it["title"])]
    filtered_total = sum(len(v) for v in filtered_grouped.values())

    st.markdown(
        f"""
        <div style='margin-top: 32px; padding: 18px 22px; background: {ICE_BLUE}; border-left: 5px solid {NAVY}; border-radius: 2px;'>
            <div style='color: {NAVY}; font-weight: 800; font-size: 22px; line-height: 1.1; margin-bottom: 4px;'>
                {generated.strftime('%B %d, %Y')}
            </div>
            <div style='color: {DARK_GREY}; font-size: 12px; letter-spacing: 0.5px; text-transform: uppercase; font-weight: 600;'>
                Past {LOOKBACK_DAYS} days &middot; {filtered_total} stories &middot; {len(filtered_grouped)} themes &middot; Source: Google News RSS
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # "This week's top news" — top 4 ranked stories
    top_picks = pick_top_stories(filtered_grouped, n=4)
    if top_picks:
        rows = []
        for i, (it, theme) in enumerate(top_picks, start=1):
            date = it["published"].strftime("%b %d")
            src = (
                f"<span style='color: rgba(255,255,255,0.55); font-style: italic; font-size: 12px;'> &mdash; {it['source']}</span>"
                if it["source"] else ""
            )
            theme_short = theme.split(" (")[0]
            rows.append(
                f"<div style='margin-bottom: 14px; line-height: 1.5;'>"
                f"<span style='color: {PERIWINKLE}; font-weight: 800; font-size: 13px; margin-right: 8px;'>{i:02d}</span>"
                f"<span style='color: rgba(255,255,255,0.55); font-size: 11px; letter-spacing: 0.8px; text-transform: uppercase; font-weight: 600; margin-right: 8px;'>{theme_short} &middot; {date}</span><br>"
                f"<a href='{it['link']}' target='_blank' style='color: #FFFFFF !important; font-size: 15.5px; font-weight: 600; text-decoration: none; border-bottom: 1px dotted rgba(255,255,255,0.4);'>{it['title']}</a>"
                f"{src}"
                f"</div>"
            )
        st.markdown(
            f"""
            <div style='margin-top: 22px; padding: 22px 26px; background: {NAVY}; border-radius: 4px;'>
                <div style='color: {PERIWINKLE}; font-size: 11px; letter-spacing: 1.8px; text-transform: uppercase; font-weight: 700; margin-bottom: 14px;'>
                    This week's top news
                </div>
                {''.join(rows)}
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Download button
    docx_bytes = build_docx_bytes(filtered_grouped, filtered_total)
    st.download_button(
        label="Download as Word doc",
        data=docx_bytes,
        file_name=f"ot_cyber_brief_{generated.strftime('%Y-%m-%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Themed sections — newspaper feel, iterates QUERIES.keys() to enforce order
    for i, theme in enumerate(QUERIES.keys()):
        items = filtered_grouped.get(theme, [])
        st.markdown(
            f"""
            <div style='margin: 36px 0 14px 0;'>
                <div style='color: {PERIWINKLE}; font-size: 11px; font-weight: 700; letter-spacing: 1.8px; text-transform: uppercase; margin-bottom: 2px;'>
                    Section {i+1:02d}
                </div>
                <h2 style='color: {NAVY}; font-weight: 800; font-size: 22px; margin: 0 0 4px 0; line-height: 1.2;'>
                    {theme}
                </h2>
                <div style='height: 2px; background: {NAVY}; width: 48px; margin-top: 8px;'></div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if not items:
            st.markdown(
                f"<div style='color: {MID_GREY}; font-style: italic; font-size: 13px; margin: 8px 0;'>No news in the lookback window.</div>",
                unsafe_allow_html=True,
            )
            continue
        story_rows = []
        for it in items:
            date = it["published"].strftime("%b %d")
            src = (
                f" <span style='color: {MID_GREY}; font-style: italic; font-size: 12px;'>&mdash; {it['source']}</span>"
                if it["source"] else ""
            )
            story_rows.append(
                f"<div style='margin-bottom: 10px; line-height: 1.45; color: #3A3A3C; font-size: 14.5px;'>"
                f"<span style='display: inline-block; min-width: 56px; color: {NAVY}; font-weight: 700; font-size: 13px;'>{date}</span>"
                f"<a href='{it['link']}' target='_blank' style='color: {NAVY};'>{it['title']}</a>"
                f"{src}"
                f"</div>"
            )
        st.markdown("".join(story_rows), unsafe_allow_html=True)

# Footer
st.markdown(
    f"""
    <div style='margin-top: 56px; padding-top: 16px; border-top: 1px solid {ICE_BLUE};'>
        <div style='color: {MID_GREY}; font-size: 11px; letter-spacing: 0.3px;'>
            (C) {datetime.now().year} Banneker Partners, LLC. All Rights Reserved. Confidential.
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)
